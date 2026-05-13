#!/usr/bin/env python3
"""
Muallimx Platform — Word documentation generator
================================================
Generates a comprehensive English .docx for developers: stack, architecture, personas,
security, commerce, live/classroom, integrations, full middleware & services catalogs,
route-by-prefix reference (public, student, employee, admin, instructor), payments,
subscription keys, config file map, data clusters, and appendix. Intended to be shared
with engineering teams as an onboarding artifact — regenerate after major refactors.

Usage (from repository root or any directory):
    pip install -r scripts/requirements-documentation.txt
    python scripts/generate_platform_documentation.py
    python scripts/generate_platform_documentation.py --output ./Muallimx_Platform_Documentation.docx

Requires: python-docx
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import subprocess
import sys
from collections import defaultdict
from datetime import timezone
from pathlib import Path

# Repository root when this file lives in <repo>/scripts/
REPO_ROOT = Path(__file__).resolve().parent.parent

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as e:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx\n"
        "Install with: pip install -r scripts/requirements-documentation.txt"
    ) from e


def _resolve_logo_png(repo: Path) -> Path | None:
    for candidate in (
        repo / "docs" / "logo.png",
        repo / "public" / "logo.png",
        repo / "public" / "images" / "logo.png",
    ):
        if candidate.is_file():
            return candidate
    return None


def _apply_header_logo_and_company(doc: Document, repo: Path) -> None:
    """
    Word header repeats on every page for each document section.
    Logo from logo.png (docs/ or public/) + implementer branding: Solvesta (short English tagline).
    """
    logo_path = _resolve_logo_png(repo)
    tagline = "Software company — implementation and delivery partner for this project."

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if logo_path is not None:
            r_pic = p.add_run()
            r_pic.add_picture(str(logo_path), width=Inches(0.9))
            p.add_run("  ")

        r_brand = p.add_run("Solvesta")
        r_brand.bold = True
        r_brand.font.size = Pt(12)
        r_brand.font.color.rgb = RGBColor(0x0F, 0x4C, 0x5C)

        p.add_run().add_break(WD_BREAK.LINE)

        r_en = p.add_run(tagline)
        r_en.font.size = Pt(9)
        r_en.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def _set_cell_shading(cell, fill_hex: str) -> None:
    """Light background for table cells (e.g. 'F8FAFC')."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _add_horizontal_line(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pb.append(bottom)
    p_pr.append(pb)


def _mono_run(paragraph, text: str, size_pt: int = 9) -> None:
    r = paragraph.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(size_pt)
    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def _bullet(doc: Document, text: str, level: int = 0) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _two_col_table(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text = header
    for c in hdr:
        _set_cell_shading(c, "EEF2FF")
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (a, b) in enumerate(rows, start=1):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b


def _chunked(seq: list, size: int):
    for i in range(0, len(seq), size):
        yield seq[i : i + size]


def _add_mono_lines(doc: Document, lines: list[str], size_pt: int = 7) -> None:
    for chunk in _chunked(lines, 50):
        p = doc.add_paragraph()
        _mono_run(p, "\n".join(chunk), size_pt)


def _try_route_list_json(repo: Path) -> list[dict] | None:
    artisan = repo / "artisan"
    if not artisan.is_file():
        return None
    try:
        proc = subprocess.run(
            ["php", str(artisan), "route:list", "--json"],
            cwd=str(repo),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300,
            check=False,
        )
        if proc.returncode != 0 or not (proc.stdout or "").strip():
            return None
        return json.loads(proc.stdout)
    except (json.JSONDecodeError, OSError, subprocess.TimeoutExpired):
        return None


def _add_auto_generated_inventory(doc: Document) -> None:
    """Populate doc from filesystem + php artisan (when available). Nothing manual — full lists."""
    repo = REPO_ROOT

    _page_break(doc)
    _heading(doc, "25. Auto-generated: Eloquent models (app/Models)", 1)
    models_dir = repo / "app" / "Models"
    models = sorted(p.name for p in models_dir.glob("*.php")) if models_dir.is_dir() else []
    doc.add_paragraph(
        f"Complete list of PHP model files ({len(models)}). Table names and relations are defined in these classes and database/migrations."
    )
    _add_mono_lines(doc, models, 7)

    _page_break(doc)
    _heading(doc, "26. Auto-generated: Artisan commands (app/Console/Commands)", 1)
    cmd_dir = repo / "app" / "Console" / "Commands"
    cmds = sorted(p.stem for p in cmd_dir.glob("*.php")) if cmd_dir.is_dir() else []
    doc.add_paragraph(
        f"Console command classes ({len(cmds)}). Register signatures in each class; schedule references in bootstrap/app.php and routes/console.php as applicable."
    )
    _add_mono_lines(doc, [c + ".php" for c in cmds], 7)

    _page_break(doc)
    _heading(doc, "27. Auto-generated: Queue jobs (app/Jobs)", 1)
    jobs_dir = repo / "app" / "Jobs"
    jobs = sorted(p.name for p in jobs_dir.glob("*.php")) if jobs_dir.is_dir() else []
    doc.add_paragraph(f"Job classes ({len(jobs)}).")
    _add_mono_lines(doc, jobs or ["(no files)"], 7)

    _page_break(doc)
    _heading(doc, "28. Auto-generated: Event listeners (app/Listeners)", 1)
    lst_dir = repo / "app" / "Listeners"
    listeners = sorted(p.name for p in lst_dir.glob("*.php")) if lst_dir.is_dir() else []
    doc.add_paragraph(f"Listener classes ({len(listeners)}).")
    _add_mono_lines(doc, listeners or ["(no files)"], 7)

    _page_break(doc)
    _heading(doc, "29. Auto-generated: HTTP controllers (complete file list)", 1)
    ctr_root = repo / "app" / "Http" / "Controllers"
    by_ns: dict[str, list[str]] = defaultdict(list)
    if ctr_root.is_dir():
        for p in ctr_root.rglob("*.php"):
            rel = p.relative_to(ctr_root)
            if len(rel.parts) == 1:
                ns = "(root)"
            else:
                ns = rel.parts[0]
            by_ns[ns].append("/".join(rel.parts))
    doc.add_paragraph(
        f"Total controller PHP files: {sum(len(v) for v in by_ns.values())}. Grouped by first subdirectory under Controllers."
    )
    for ns in sorted(by_ns.keys(), key=lambda x: (x == "(root)", x.lower())):
        lines = [f"=== {ns} ==="] + sorted(by_ns[ns])
        _add_mono_lines(doc, lines, 6)
        doc.add_paragraph()

    _page_break(doc)
    _heading(doc, "30. Auto-generated: Laravel routes (php artisan route:list --json)", 1)
    routes = _try_route_list_json(repo)
    if routes is None:
        doc.add_paragraph(
            "Could not run `php artisan route:list --json` (PHP not in PATH, wrong cwd, or artisan failure). "
            "Run from the project root on a machine with PHP and regenerate this document."
        )
    else:
        doc.add_paragraph(
            f"Total registered routes: {len(routes)}. Each line: METHOD | URI | route name | Controller@method"
        )
        lines_out: list[str] = []

        def _sort_key(r: dict):
            return (r.get("uri") or "", r.get("method") or "", r.get("name") or "")

        for r in sorted(routes, key=_sort_key):
            m = (r.get("method") or "").replace("|", "/")
            u = r.get("uri") or ""
            n = r.get("name") or ""
            act = r.get("action") or ""
            if isinstance(act, str) and "@" in act:
                action_s = act.split("App\\Http\\Controllers\\", 1)[-1] if "App\\Http\\Controllers\\" in act else act
            else:
                action_s = str(act)[:80]
            lines_out.append(f"{m:14} | {u:52} | {n:45} | {action_s}")
        _add_mono_lines(doc, lines_out, 5)

    _page_break(doc)
    _heading(doc, "31. Auto-generated: Database migrations", 1)
    mig_dir = repo / "database" / "migrations"
    migs = sorted(p.name for p in mig_dir.glob("*.php")) if mig_dir.is_dir() else []
    doc.add_paragraph(
        f"Total migration files: {len(migs)}. Schema history for every table change; pair with app/Models."
    )
    _add_mono_lines(doc, migs, 6)

    _page_break(doc)
    _heading(doc, "32. Auto-generated: Configuration files (config/)", 1)
    cfg_dir = repo / "config"
    cfgs = sorted(p.name for p in cfg_dir.glob("*.php")) if cfg_dir.is_dir() else []
    doc.add_paragraph(f"Total config PHP files: {len(cfgs)}.")
    _add_mono_lines(doc, cfgs, 7)

    _page_break(doc)
    _heading(doc, "33. Auto-generated: Language files (lang/)", 1)
    lang_dir = repo / "lang"
    lang_files = sorted(p.relative_to(lang_dir).as_posix() for p in lang_dir.rglob("*.php")) if lang_dir.is_dir() else []
    doc.add_paragraph(f"Total PHP translation / lang files: {len(lang_files)}.")
    _add_mono_lines(doc, lang_files or ["(no files)"], 7)

    _page_break(doc)
    _heading(doc, "34. Auto-generated: Service providers (app/Providers)", 1)
    prov_dir = repo / "app" / "Providers"
    provs = sorted(p.name for p in prov_dir.glob("*.php")) if prov_dir.is_dir() else []
    _add_mono_lines(doc, provs or ["(no files)"], 7)

    _page_break(doc)
    _heading(doc, "36. Auto-generated: Database seeders (database/seeders)", 1)
    seed_dir = repo / "database" / "seeders"
    seeds = sorted(p.name for p in seed_dir.glob("*.php")) if seed_dir.is_dir() else []
    doc.add_paragraph(f"Seeder classes ({len(seeds)}). Run via php artisan db:seed --class=...")
    _add_mono_lines(doc, seeds or ["(no files)"], 7)


def _add_engineering_deep_dive(doc: Document) -> None:
    """
    Extended reference for engineers inheriting the codebase.
    Sourced from routes/web.php, app/Http/Middleware, app/Services, config/.
    """
    _heading(doc, "10. Repository layout (read this first)", 1)
    doc.add_paragraph(
        "Almost all HTTP routes live in routes/web.php (there is no separate routes/api.php in this "
        "project). Controllers are namespaced under app/Http/Controllers. Blade views under "
        "resources/views; translations under lang/. Database under database/migrations and "
        "database/seeders."
    )
    _two_col_table(
        doc,
        [
            ("app/Http/Controllers/Admin", "Back-office: users, finance, courses, HR, live ops, marketing."),
            ("app/Http/Controllers/Instructor", "Teacher dashboard: curriculum, lectures, exams, live, payouts."),
            ("app/Http/Controllers/Student", "Learner: orders, wallet, classroom, portfolio, AI usages."),
            ("app/Http/Controllers/Employee", "Internal desks: sales, HR, recruitment, tasks, accounting."),
            ("app/Http/Controllers/Public", "Marketing site, catalog, checkout, certificate verify."),
            ("app/Http/Controllers/Community", "Contributor registration and community UI."),
            ("app/Http/Controllers/Auth", "Login, register, password reset, 2FA challenge."),
            ("app/Http/Controllers/Api", "Webhooks and n8n callbacks (not a full REST API surface)."),
            ("app/Services", "Payments, checkout pricing, Fawaterak/Kashier, subscriptions, AI client, uploads."),
            ("app/Http/Middleware", "RBAC, ownership, video protection, locale, concurrent sessions."),
            ("resources/views/layouts", "student-sidebar.blade.php, app layout, admin shell."),
            ("config/", "student_subscription_features, rbac_*, fawaterak, kashier, security, filesystems."),
        ],
        ("Location", "Purpose"),
    )
    _page_break(doc)

    _heading(doc, "11. Middleware catalog (app/Http/Middleware)", 1)
    doc.add_paragraph(
        "Register aliases in bootstrap/app.php or app/Http/Kernel equivalent (Laravel 11+ style). "
        "These gates wrap controllers; when debugging 403/redirect, check the route’s middleware stack."
    )
    _two_col_table(
        doc,
        [
            ("CheckActiveStatus", "Blocks inactive users from proceeding after login."),
            ("CheckRole", "role:student|instructor|teacher style role checks."),
            ("EnsurePermission", "permission:admin.access and granular keys; super_admin bypass patterns."),
            ("EnsureOwnership", "ownership:order,order — user must own the model instance."),
            ("EnsureGuestOnly", "guest-only for registration flows that must not run for logged-in users."),
            ("EnsureTwoFactorEnabled", "Forces 2FA where platform policy requires it (available; may be commented in bootstrap)."),
            ("EnsureStudentAiUsagesAccess", "student.ai-usages — package + permission for AI usage UI."),
            ("EnsureEmployeeCan", "employee.can:sales_desk — job JSON permissions on EmployeeJob."),
            ("EnsureCommunityContributor", "Restricts contributor-only routes (alias community.contributor)."),
            ("PreventConcurrentSessions", "prevent-concurrent — blocks parallel sessions for same account."),
            ("RestrictRbacEmployeeAdminRoutes", "rbac.strict.admin — limits admin UI for RBAC employees."),
            ("VideoProtectionMiddleware", "Guards lesson/exam video delivery endpoints."),
            ("SetLocale", "Appends to web group: language from session/query."),
            ("SetLandingLocale", "Alias landing.locale for localized landing."),
            ("SecurityHeadersMiddleware", "Prepends globally: security-related HTTP headers."),
            ("InputSanitizationMiddleware", "Web group: sanitizes dangerous input patterns."),
            ("FileUploadSecurityMiddleware", "Web group: upload validation."),
            ("LogActivityMiddleware", "Appends globally: activity / audit logging."),
            ("EnhancedRateLimitingMiddleware", "Named rate limit profiles where referenced."),
        ],
        ("Middleware class (19 files in app/Http/Middleware)", "Role"),
    )
    _heading(doc, "11.1 Bootstrap middleware & CSRF (bootstrap/app.php)", 2)
    for line in (
        "Global: SecurityHeadersMiddleware; LogActivityMiddleware on all requests.",
        "Web group appended: SetLocale, InputSanitizationMiddleware, FileUploadSecurityMiddleware, CheckActiveStatus.",
        "CSRF validation excludes: api/n8n/* and api/live-recordings/register (external webhooks).",
        "Middleware aliases: role, permission, ownership, guest-only, prevent-concurrent, landing.locale, community.contributor, employee.can, rbac.strict.admin, student.ai-usages.",
        "AuthenticationException sends guests under /community/* to route('community.login') — verify route:list; see section 23.",
    ):
        _bullet(doc, line)
    _page_break(doc)

    _heading(doc, "12. Domain services — complete list (app/Services)", 1)
    doc.add_paragraph(
        "Every PHP class under app/Services/ (root and Community/) is listed below — one row per file. "
        "This is the full service layer inventory for the platform as shipped in the repository."
    )
    _two_col_table(
        doc,
        [
            ("AdminPanelBranding.php", "Admin UI branding / theme-related settings."),
            ("AssignmentFileStorage.php", "Secure storage paths for assignment submission files."),
            ("CalendarNotificationService.php", "Calendar-driven reminders / notifications."),
            ("ClassroomSubscriptionFeatureMenuService.php", "Maps subscription features to classroom menu entries."),
            ("ContactMessageAlertService.php", "Alerts staff when contact form messages arrive."),
            ("CouponCommissionService.php", "Coupon commission accruals, expenses, payouts logic."),
            ("CourseCheckoutPricingService.php", "Checkout quote: price, coupon, wallet, installments."),
            ("CurriculumLibraryR2MultipartService.php", "Multipart / presigned uploads for curriculum library on object storage."),
            ("EmailNotificationService.php", "Sends email broadcast campaigns."),
            ("ExcelExportService.php", "Spreadsheet exports for admin reports."),
            ("FawaterakApiService.php", "Low-level Fawaterak HTTP API client."),
            ("FawaterakService.php", "High-level Fawaterak payment session and return handling."),
            ("FullAiSuiteContextService.php", "Builds context payloads for full AI suite features."),
            ("InstructorCoursePercentageService.php", "Instructor revenue share percentages per course/rules."),
            ("InstructorMarketingRankingService.php", "Ranking / scoring for instructor marketing surfaces."),
            ("KashierService.php", "Kashier redirect, callback, and payment status handling."),
            ("MuallimxAiClient.php", "HTTP client to Muallimx AI backend (config/muallimx_ai.php)."),
            ("OrderWalletAndCouponFinalizer.php", "Post-payment order finalization with wallet and coupons."),
            ("PaymentGatewaySettings.php", "Reads gateway credentials / modes from settings storage."),
            ("PlatformCourseCertificateService.php", "Generates course completion certificates (PDF pipeline)."),
            ("PlatformSecuritySettings.php", "Platform-wide security toggles and secrets resolution."),
            ("PortfolioImageStorage.php", "Portfolio project image storage on disk/S3."),
            ("PublicFooterSettings.php", "Footer links / CMS blocks for public layout."),
            ("ReferralService.php", "Referral codes, attribution, rewards."),
            ("SecurityService.php", "Security helpers (password policy, lockouts, etc. — inspect class)."),
            ("ServerSshService.php", "SSH connectivity for admin live-server file browser."),
            ("SiteServiceImageStorage.php", "Images for SiteService CMS pages."),
            ("SiteTestimonialImageStorage.php", "Testimonial avatars / media."),
            ("StatisticsCacheService.php", "Caches expensive admin statistics queries."),
            ("SubscriptionLimitService.php", "Enforces plan limits (courses, storage, features)."),
            ("SupportTicketAlertService.php", "Notifies staff on new or escalated support tickets."),
            ("TeacherSubscriptionActivationService.php", "Activates teacher subscriptions after gateway or admin action."),
            ("TeamsAttendanceImportService.php", "Parses Teams attendance exports into lecture attendance."),
            ("UserProfileImageStorage.php", "User profile avatar storage."),
            ("WhatsAppService.php", "WhatsApp Cloud / API sends for transactional messaging."),
            ("Community/CommunityRegistrationService.php", "Registers and onboards community contributors."),
            ("Community/DatasetFileReaderService.php", "Reads/parses uploaded dataset files for community features."),
        ],
        ("Service file", "Responsibility (summary)"),
    )
    _heading(doc, "12.1 Composer autoload helpers (app/Helpers)", 2)
    doc.add_paragraph(
        "Non-class helpers loaded via composer.json \"files\": FilesystemHelper.php, WhatsAppHelper.php — "
        "use for path normalization and WhatsApp formatting without injecting a service container."
    )
    _heading(doc, "12.2 Scheduled tasks (bootstrap/app.php)", 2)
    for line in (
        "reports:send-monthly — monthly on day 1 at 09:00.",
        "Monthly cleanup: deletes old WhatsAppMessage and ActivityLog rows.",
        "Daily cache of active_users_today from ActivityLog.",
        "installments:process — daily 08:00.",
        "subscriptions:expire — daily 00:05.",
        "live:send-reminders --minutes=10 — every minute (live session reminders).",
        "live:auto-end-sessions — every five minutes.",
    ):
        _bullet(doc, line)
    _page_break(doc)

    _heading(doc, "13. Public (guest) surface — marketing & commerce entry", 1)
    for line in (
        "GET / — LandingController (localized via SetLandingLocale where applied).",
        "Static pages: /about, /faq, /terms, /privacy, /pricing, /team, /certificates, /help, /refund, /testimonials, /events, /partners.",
        "Portfolio marketing: /portfolio, /portfolio/teacher/{id}, /portfolio/{id}.",
        "Services CMS: /services, /services/{slug}.",
        "Media blog: /media, /media/{id}.",
        "Catalog: /courses (filtering), /instructors, /instructors/{id}, /course/{id} (course landing).",
        "Checkout: /course/{courseId}/checkout (+ quote, complete, Kashier, Fawaterak prepare/methods/pay), /checkout/kashier/callback, /checkout/fawaterak/{status}.",
        "Teacher SaaS pricing: /pricing, /pricing/checkout/{plan}, Fawaterak endpoints under same area.",
        "Certificate verification: /certificates/verify and /certificates/verify/{code}.",
        "Classroom guest join: /classroom/join/{code} (+ enter, heartbeat, leave, share-annotation).",
        "Contact: GET/POST /contact.",
        "SEO: GET /sitemap.xml builds URL set from AdvancedCourse, instructors, Media, SiteService.",
        "Assets: GET /storage/{path} serves storage/app/public with path traversal protection; GET /mx-vendor/excalidraw/{path} serves whiteboard JS/CSS with correct MIME when public/vendor is not web-root.",
        "Guest auth: /login, /register, /forgot-password, /reset-password/*, /2fa/challenge — throttled.",
        "Legacy compatibility: GET /package/{slug} (active Package with related courses); learning-path URLs redirect to /courses or stubs; GET /learning-paths → 301 /courses.",
        "Fawaterak embed: 301 /fawaterk/plugin.min.js → /js/checkout-pay-widget.v1.js (FawaterkPluginController, throttled, some web middleware disabled for ad-blocker compatibility).",
    ):
        _bullet(doc, line)
    _page_break(doc)

    _heading(doc, "14. Authenticated shell — middleware group [auth, prevent-concurrent]", 1)
    doc.add_paragraph(
        "Everything in this group requires a logged-in user and uses concurrent-session protection. "
        "DashboardController@GET /dashboard is the persona router."
    )
    _heading(doc, "14.1 Student-facing paths at URL root (not under /student)", 2)
    for line in (
        "Browse catalog while logged in: /academic-years, /academic-years/{year}/subjects, /subjects/{subject}/courses, /courses/{advancedCourse}.",
        "role:student — /my-courses*, /orders*, /exams/* (student.exams.*), /profile, /settings, /notifications*, /calendar + /api/calendar/events.",
        "Consultations: /consultations* (request paid sessions with instructors).",
        "Portfolio: /my-portfolio/* (projects + marketing profile CRUD).",
        "Subscription: GET /my-subscription (student.my-subscription).",
        "AI usages: /ai-usages* behind student.ai-usages middleware (package + permission).",
        "Muallimx Classroom (student as host): /classroom* CRUD, room, recordings presign/upload/complete, audio tracks, AI report per meeting.",
        "Referrals: /referrals* for students.",
        "POST /api/validate-coupon, POST /api/video/info, GET /api/lessons/{lesson} (JSON for enrolled student).",
        "GET /api/courses/{course}/students — JSON list for instructor|teacher only (instructor grading tools).",
        "POST /courses/{advancedCourse}/order — create order from student.",
    ):
        _bullet(doc, line)

    _heading(doc, "14.2 Shared student|instructor|teacher routes", 2)
    for line in (
        "/support* — SupportTicketController (tickets and replies).",
        "/features/{feature} — SubscriptionFeatureController; route constraint whitelist: library_access, ai_tools, classroom_access, support, visible_to_academies, can_apply_opportunities, full_ai_suite, teacher_evaluation, recommended_to_academies, priority_opportunities, direct_support.",
        "POST /features/full-ai-suite/preview — throttled AI preview.",
        "/curriculum-library* — index, item by slug, material download/view (html/pdf/presentation), file download/view.",
        "/academies/visibility + POST apply to academy opportunities.",
        "/teaching-opportunities* — listing and apply.",
    ):
        _bullet(doc, line)

    _heading(doc, "14.3 Prefix /student (role:student) — invoices, wallet, certificates, assignments, tasks, live", 2)
    for line in (
        "Resource-style: student.invoices.*, student.wallet.*, student.certificates.* (+ file), student.achievements.*.",
        "student.assignments.* with ownership middleware; presigned direct-to-storage submission uploads (throttled).",
        "student.tasks.* — assigned operational tasks visible to the student.",
        "student.live-sessions.* — join/leave/status/share-annotation for scheduled live sessions.",
        "student.live-recordings.index/show — published recordings only (comment in routes references R2).",
    ):
        _bullet(doc, line)
    _page_break(doc)

    _heading(doc, "15. Prefix /employee — internal operations (EnsureEmployeeCan)", 1)
    doc.add_paragraph(
        "Each sub-route checks employee.can:{capability} where capability is stored on the employee’s job "
        "record (JSON permissions). Typical capabilities used in routes:"
    )
    _two_col_table(
        doc,
        [
            ("dashboard", "GET /employee/dashboard"),
            ("desk_accountant", "GET /employee/desk/accountant"),
            ("sales_desk", "/employee/sales/* — desk, orders, notes, claim, leads CRUD, convert, lost"),
            ("hr_desk", "/employee/hr/* — leaves approve/reject, employee directory, HR events, recruitment (openings, candidates, applications, interviews)"),
            ("supervision_desk", "GET /employee/desk/supervision"),
            ("academic_supervision_desk", "/employee/desk/academic-supervision/* — observe classroom meetings"),
            ("tasks", "/employee/tasks* — status updates and deliverables"),
            ("leaves", "/employee/leaves* — employee self-service leave requests"),
            ("accounting", "/employee/accounting — bank account maintenance"),
            ("agreements", "/employee/agreements* — view employment agreements"),
            ("profile", "/employee/profile"),
            ("notifications", "/employee/notifications* + polling endpoints"),
            ("calendar", "/employee/calendar + events API"),
            ("reports", "/employee/reports"),
            ("settings", "/employee/settings"),
        ],
        ("employee.can key", "Surface"),
    )
    _page_break(doc)

    _heading(doc, "16. Prefix /admin — back-office (permission:admin.access + rbac.strict.admin)", 1)
    doc.add_paragraph(
        "Large surface area; below is a grouped map. Many routes use throttle:* for destructive or "
        "expensive operations. Sub-routes may require extra permissions (e.g. users.permissions, "
        "academic_supervision.manage)."
    )
    _two_col_table(
        doc,
        [
            ("Core", "dashboard, profile, nav-notifications poll, users CRUD, students-accounts, activity-log, two-factor-logs, statistics"),
            ("Student control", "students-control/paid-features, consumption per user"),
            ("Portfolio marketing review", "portfolio-marketing-profiles approve/reject"),
            ("Catalog", "course-categories, advanced-courses (+ students, orders, stats, duplicate, activate-student), courses/{course}/lessons, legacy subjects/courses resources"),
            ("Assessments", "question-bank, question-categories, exams (+ questions, publish, statistics, preview)"),
            ("Lectures & assignments", "lectures, assignments (+ submissions, grade), tasks admin"),
            ("Portfolio QC", "admin portfolio index/show approve/reject/publish/toggle-visibility"),
            ("CMS / site", "about, contact-messages, faq, site-services, site-testimonials"),
            ("System", "system-settings (+ platform 2FA enable/disable flow), packages/pricing, video-providers"),
            ("Notifications", "notifications*, employee-notifications*, email-broadcasts/{audience}"),
            ("Enrollments", "online-enrollments* (quick-activate, progress, search-by-phone)"),
            ("HR / workforce", "employees*, employee-jobs*, employee-tasks*, leaves*, academic-supervision*"),
            ("Instructor ops", "instructor-requests*, quality-control*, personal-branding*, withdrawals*, instructor-accounts"),
            ("Finance", "invoices*, payments*, transactions*, wallets* (duplicate resource blocks with throttles), expenses*, accounting/reports*, salaries/instructor payouts"),
            ("Subscriptions", "subscriptions*, subscription-requests approve/reject, teacher-features, consumption"),
            ("Support & consultations", "support-tickets*, support-inquiry-categories, consultations admin"),
            ("Hiring", "hiring-academies*, academy-opportunities* + applications + recruitment desk presentations"),
            ("Curriculum library admin", "curriculum-library categories/items/structure/materials with S3 multipart/presign endpoints"),
            ("Marketing", "popup-ads, coupons, coupon-commissions, referral-programs, referrals, loyalty, student-wallet-credit"),
            ("Certificates & engagement", "certificates* (+ design, preview, file), achievements*, badges*, reviews*"),
            ("Attendance & performance", "attendance*, performance (cache/temp/db tools)"),
            ("Live stack", "live-sessions*, live-servers* (SSH browse), live-recordings*, classroom-recordings*, live-settings*"),
            ("Automation", "n8n live-session-reports + settings/test-connection"),
            ("Installments & contracts", "installments/plans+agreements, agreements (instructor), employee-agreements"),
            ("Sales (admin view)", "sales/leads, sales analytics"),
            ("Messages / WhatsApp", "messages/* including WhatsApp API settings, templates, monthly reports"),
        ],
        ("Admin area", "What lives there (high level)"),
    )
    doc.add_paragraph(
        "Note: Several legacy admin paths for academic-years, learning-paths, academic-subjects, and "
        "learning-path-enrollments are explicitly redirected to advanced-courses or online-enrollments "
        "to avoid broken bookmarks while the product consolidated on AdvancedCourse + online enrollments."
    )
    _page_break(doc)

    _heading(doc, "17. Prefix /instructor — teacher dashboard (role:instructor|teacher)", 1)
    for line in (
        "Calendar + API events; consultations index/show.",
        "Reuses Student\\ClassroomController under /instructor/classroom/* for meeting host flows (room, recordings, AI report).",
        "Profile + personal-branding submit for admin review.",
        "courses.index/show; courses/{course}/curriculum — sections, curriculum items, exams/assignments from curriculum, lecture video questions.",
        "Legacy /instructor/courses/{course}/lessons/* redirects to curriculum or lectures index (lessons system retired).",
        "lectures resource + Teams attendance sync + manual attendance update; learning-path routes redirect to dashboard.",
        "assignments + submissions + grade; exams + exam questions from bank or new; question-banks + questions.",
        "attendance index + per lecture; tasks + deliverables + progress.",
        "management-requests CRUD-lite for escalations to admin.",
        "agreements index/show + export activations; transfer-account; withdrawals create/list/show/cancel.",
        "live-sessions: create, start, room, student-whiteboard, share-annotations, audio presign/complete, ai-report, end.",
    ):
        _bullet(doc, line)
    _page_break(doc)

    _heading(doc, "18. Payments, orders, and gateways (technical flow)", 1)
    p = doc.add_paragraph()
    _mono_run(
        p,
        "Course purchase:\n"
        "  GET  /course/{id}/checkout\n"
        "  POST /course/{id}/checkout/quote          — CourseCheckoutPricingService\n"
        "  POST /course/{id}/checkout/complete\n"
        "  POST /course/{id}/checkout/kashier      — Kashier redirect\n"
        "  POST /course/{id}/checkout/fawaterak/*  — prepare, methods, pay\n"
        "  GET  /checkout/kashier/callback\n"
        "  GET  /checkout/fawaterak/{status}       — return; may correlate subscription/order\n"
        "  POST /course/{id}/enroll-free\n"
        "Teacher subscription (public):\n"
        "  /pricing/checkout/{plan} + store + fawaterak/*\n"
        "Webhooks / misc API in web.php (no CSRF per bootstrap exception where noted):\n"
        "  POST /api/live-recordings/register\n"
        "  PATCH|POST /api/n8n/live-session-reports/{report}\n"
        "  PATCH|POST /api/n8n/classroom-meeting-reports/{report}\n",
        8,
    )
    doc.add_paragraph(
        "Wallet balances live on Wallet / WalletTransaction models; admin can credit marketing wallets. "
        "Coupons validated via /api/validate-coupon during checkout UI."
    )
    _page_break(doc)

    _heading(doc, "19. Subscription entitlements & config keys", 1)
    doc.add_paragraph(
        "Teacher-side subscription rows carry JSON feature flags. Student UI merges activeSubscription(), "
        "hasSubscriptionFeature($key), and sometimes explicit permissions (e.g. AI usages). "
        "config/student_subscription_features.php binds feature_key → named route + icons for sidebar."
    )
    _two_col_table(
        doc,
        [
            ("library_access", "Curriculum library routes."),
            ("ai_tools", "student.features.show feature=ai_tools."),
            ("classroom_access", "Muallimx Classroom host routes."),
            ("support", "Support tickets."),
            ("teacher_profile", "Portfolio + admin portfolio-marketing-profiles review."),
            ("visible_to_academies", "Academy visibility page."),
            ("can_apply_opportunities", "Teaching opportunities apply."),
            ("full_ai_suite", "Full AI suite pages + throttled preview endpoint."),
            ("teacher_evaluation", "Feature page placeholder / content."),
            ("recommended_to_academies", "Feature page."),
            ("priority_opportunities", "Feature page."),
            ("direct_support", "Routes to support (priority channel branding)."),
        ],
        ("feature_key", "Typical UX"),
    )
    doc.add_paragraph(
        "Admin: StudentControlController uses the same keys for paid-features and per-user consumption reports."
    )
    _page_break(doc)

    _heading(doc, "20. Configuration files every engineer should open", 1)
    _two_col_table(
        doc,
        [
            ("config/student_subscription_features.php", "Maps subscription keys to student routes and admin review routes."),
            ("config/rbac_admin_route_access.php", "Which admin paths each RBAC role may hit."),
            ("config/rbac_permission_sidebar.php", "Sidebar visibility by permission."),
            ("config/admin_sidebar_role_map.php", "Legacy/admin menu mapping."),
            ("config/employee_sidebar.php", "Employee navigation structure."),
            ("config/fawaterak.php / kashier.php", "Gateway keys and modes."),
            ("config/filesystems.php", "Disks: local, public, s3, possibly R2-compatible."),
            ("config/security.php / upload_limits.php", "Hardening and upload caps."),
            ("config/muallimx_ai.php", "AI endpoint keys and defaults."),
            ("config/certificates.php", "Certificate layout and issuance rules."),
            ("config/permission_aliases.php", "Compatibility aliases for renamed permissions."),
        ],
        ("File", "Why open it"),
    )
    _page_break(doc)

    _heading(doc, "21. Data model — additional clusters (ORM)", 1)
    _two_col_table(
        doc,
        [
            ("CurriculumLibrary*", "Categories, items, sections, materials, files — interactive curriculum product."),
            ("AcademyOpportunity*", "Hiring academies post jobs; students apply; admin recruitment desk."),
            ("Consultation* / ConsultationSetting", "Paid 1:1 instructor consultations, payment confirmation workflow."),
            ("SalesLead / SalesOrderNote", "CRM-style pipeline for sales employees and admin."),
            ("LiveServer / LiveSetting / LiveSessionReport", "Operational live streaming infrastructure."),
            ("IntegrationSetting / VideoProvider", "Third-party video and integration credentials."),
            ("PopupAd / SiteService / FAQ / Media", "Marketing content models."),
            ("EmployeeAgreement / Installment*", "Contracts and staged payments for staff and students."),
        ],
        ("Cluster", "Notes"),
    )
    _page_break(doc)

    _heading(doc, "22. Tasks route at root", 1)
    doc.add_paragraph(
        "Route::resource('tasks', TaskController::class) is registered for all authenticated users inside "
        "the auth group — separate from admin employee tasks and instructor/student task controllers. "
        "Verify which Task model/policies apply when extending."
    )
    _page_break(doc)

    _heading(doc, "23. Community module (code vs routes)", 1)
    doc.add_paragraph(
        "The repository still contains Community controllers (e.g. Community/AuthController, "
        "ContributorController, CommunityPageController), views under resources/views/community and "
        "resources/views/public/community, Admin CommunityController, middleware alias "
        "community.contributor, and bootstrap exception handling that redirects unauthenticated "
        "users on /community/* to route('community.login')."
    )
    doc.add_paragraph(
        "Important: as of the current routes/web.php in this codebase, php artisan route:list --path=community "
        "returns no routes — the /community route group is not registered in web.php. "
        "Blade files that call route('community.login') may break until routes are re-wired or views updated. "
        "Engineers must run route:list after pull and either restore community routes or remove dead links."
    )
    _page_break(doc)

    _heading(doc, "24. How to stay accurate after code changes", 1)
    for line in (
        "Run: php artisan route:list --columns=Method,URI,Name,Middleware > routes-export.txt and diff after refactors.",
        "Regenerate this Word file whenever major route or permission changes ship.",
        "Use IDE 'Find usages' on route() names in Blade to trace UI → controller.",
        "Migrations are authoritative for columns; models casts() for JSON/array fields.",
    ):
        _bullet(doc, line)


def build_document(repo_hint: str = "") -> Document:
    doc = Document()
    _apply_header_logo_and_company(doc, REPO_ROOT)
    core = doc.core_properties
    core.title = "Muallimx Platform — Technical & Product Documentation"
    core.subject = "Source orientation, dashboards, business logic"
    core.keywords = "Laravel, LMS, EdTech, RBAC, subscriptions, payments"
    core.category = "Engineering Documentation"
    core.author = "Muallimx Documentation Generator"
    core.comments = "Generated by scripts/generate_platform_documentation.py"
    core.created = dt.datetime.now(timezone.utc)

    # --- Title block ---
    title = doc.add_heading("Muallimx Platform", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Technical & Business Documentation\n"
        "For engineering teams, architects, and product stakeholders"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Document generated: {dt.datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n"
        f"Repository path (hint): {repo_hint or '(run from project root)'}"
    )
    mr.font.size = Pt(10)
    mr.font.italic = True
    mr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "This document describes the Muallimx web application: a large Laravel-based "
        "learning-management and operations platform combining public marketing pages, "
        "student learning, instructor authoring, administrative control, internal employee "
        "desks (HR / sales / recruitment), community data contributions, live classrooms, "
        "subscriptions, and multi-gateway commerce. Use it to navigate the codebase, "
        "understand role-based flows, and align product behavior with implementation. "
        "Section 12 enumerates every file in app/Services/; section 11 lists every middleware class; "
        "sections 25–36 are regenerated from disk and `php artisan route:list --json` (models, commands, "
        "jobs, listeners, every controller file, every route name, every migration, every config/*.php, "
        "every lang/*.php, providers, seeders); "
        "section 23 records community route drift — always verify with route:list on your branch."
    )
    ir.font.size = Pt(11)

    _add_horizontal_line(doc.add_paragraph())

    _heading(doc, "0. Business domain summary", 1)
    doc.add_paragraph(
        "Muallimx is an integrated education platform: it sells and delivers online courses, "
        "manages instructors and internal staff, operates marketing and support channels, "
        "and monetizes through course sales, wallets, installments, and instructor subscription "
        "plans that unlock premium tooling and visibility for educators. The same codebase "
        "hosts student learning, live teaching, assessments, certificates, and operational "
        "back-office (HR, sales, recruitment) for the organization running the deployment."
    )

    # --- 1. Stack ---
    _heading(doc, "1. Technology stack", 1)
    _two_col_table(
        doc,
        [
            ("Runtime", "PHP 8.2+"),
            ("Framework", "Laravel 12"),
            ("Primary HTTP routes", "routes/web.php (single large route file; API-style endpoints embedded where needed)"),
            ("Auth", "Session-based auth, optional Google (Socialite), TOTP 2FA (pragmarx/google2fa)"),
            ("Storage", "Local + AWS S3 (league/flysystem-aws-s3-v3)"),
            ("Documents / exports", "mPDF, PhpSpreadsheet, PhpWord (server-side PDF/Excel/Word)"),
            ("Front-end", "Blade templates, Tailwind-oriented utility classes in views, Font Awesome icons"),
        ],
        ("Layer", "Details"),
    )
    doc.add_paragraph()

    # --- 2. High-level architecture ---
    _heading(doc, "2. Architectural overview", 1)
    doc.add_paragraph(
        "The application follows classic Laravel MVC: HTTP controllers under "
        "app/Http/Controllers organized by namespace (Admin, Instructor, Student, Employee, "
        "Public, Community, Auth, Api). Views live in resources/views. Domain logic is often "
        "extracted to app/Services for checkout, subscriptions, wallet finalization, classroom "
        "features, etc. Eloquent models in app/Models map to relational tables."
    )
    _bullet(doc, "Entry dashboard: DashboardController@index resolves the authenticated user’s persona and redirects or renders the correct home surface.")
    _bullet(doc, "Authorization combines legacy role string on users.role with granular permissions (permission:* middleware) and Spatie-style role relations for employees.")
    _bullet(doc, "Heavy route surface: most product URLs are declared in routes/web.php; grep route names to locate controllers.")
    _heading(doc, "2.1 Request lifecycle (simplified)", 2)
    p = doc.add_paragraph()
    _mono_run(
        p,
        "HTTP Request → web.php (middleware: auth, role, permission, throttle, …) "
        "→ Controller → Service (optional) → Eloquent → Blade view / JSON / redirect",
        9,
    )

    _page_break(doc)

    # --- 3. Personas & dashboards ---
    _heading(doc, "3. User personas and dashboards", 1)
    doc.add_paragraph(
        "After login, DashboardController inspects is_active, employee flags, role string, "
        "and permission-backed admin access. Invalid or missing roles force logout with a message."
    )
    _two_col_table(
        doc,
        [
            ("super_admin / admin", "Redirect to admin.dashboard (full admin shell with permission gates)."),
            ("instructor / teacher", "Instructor dashboard: courses, lessons, exams, live sessions, calendar, agreements, withdrawals, management requests, etc."),
            ("student", "Student dashboard: progress, courses, assignments, certificates, wallet, subscription card, AI usages (when entitled)."),
            ("Employee (is_employee)", "If linked to RBAC roles → admin dashboard with scoped permissions; else employee.dashboard (internal HR / sales / recruitment tools)."),
        ],
        ("Persona", "Primary destination"),
    )

    _heading(doc, "3.1 Student experience (high level)", 2)
    for line in (
        "Browse and purchase courses (public catalog, checkout with coupons, wallet, installments where applicable).",
        "My Courses: progress, lectures, video tokens, assignments, points.",
        "Exams, certificates, calendar, notifications, invoices, orders.",
        "Paid subscription features (from teacher package / subscription JSON): curriculum library, AI tools, virtual classroom, support tickets, portfolio, academy visibility, teaching opportunities, full AI suite, etc. — see config/student_subscription_features.php.",
        "Live sessions listing, live recordings, classroom meetings (host/join, whiteboard, recordings, AI session reports).",
        "Student AI usages index with saved games (gated by package + permission).",
    ):
        _bullet(doc, line)

    _heading(doc, "3.2 Instructor / teacher experience", 2)
    for line in (
        "Course and lesson management (including advanced courses), question banks, exams, attendance.",
        "Live sessions: scheduling, going live, student whiteboard sync, reporting.",
        "Learning paths, curriculum tools, profile and branding, management requests to admin.",
        "Financial: instructor agreements, payout / withdrawal requests.",
        "Calendar integration for teaching schedule.",
    ):
        _bullet(doc, line)

    _heading(doc, "3.3 Admin experience", 2)
    for line in (
        "Users, roles, permissions, student accounts, student control (paid features, consumption).",
        "Catalog: subjects, academic years, course categories, courses, lessons, reviews, advanced courses.",
        "Commerce: orders, invoices, transactions, packages, coupons, referrals, loyalty, installments, agreements.",
        "Operations: notifications, email broadcasts, support tickets, contact messages, site content (testimonials, services, popups, media).",
        "HR & workforce: employees, jobs, salaries, tasks, leave, performance, attendance, recruitment desk, hiring academies.",
        "Live infrastructure: live servers, live sessions admin, live recordings, classroom recordings, live settings, n8n report hooks.",
        "Quality control, academic supervision, instructor requests, instructor accounts, withdrawals.",
        "System settings, integrations (e.g. video providers, WhatsApp settings), 2FA audit logs.",
        "Community datasets, models, competitions; portfolio marketing review.",
    ):
        _bullet(doc, line)

    _heading(doc, "3.4 Employee desks", 2)
    for line in (
        "Prefix /employee with sub-areas: sales_desk (leads, orders notes), hr_desk (directory, applications, interviews), recruitment pipelines.",
        "Access controlled via employee.can:* style middleware and employee job linkage on users.",
    ):
        _bullet(doc, line)

    _page_break(doc)

    # --- 4. Security & RBAC ---
    _heading(doc, "4. Security, authentication, and authorization", 1)
    _two_col_table(
        doc,
        [
            ("Session auth", "Standard Laravel guards; guest routes for login/register/password reset."),
            ("2FA", "Routes under /2fa for setup, enable, disable; challenge after login when enabled."),
            ("Throttling", "Login and password reset routes use rate limiting middleware."),
            ("Permissions", "Granular keys such as student.view.*, admin.access, academic_supervision.manage — enforced in routes and Blade @hasPermission."),
            ("Strict admin", "rbac.strict.admin middleware tightens admin surface for non-super users."),
            ("Concurrent sessions", "prevent-concurrent middleware on selected authenticated groups."),
        ],
        ("Mechanism", "Purpose"),
    )
    doc.add_paragraph(
        "When extending the system, register new permissions in seeders (e.g. PermissionsSeeder, "
        "PermissionsAndRolesSeeder) and assign to roles; mirror checks in User model helpers where "
        "business rules depend on subscription features or legacy columns."
    )

    _page_break(doc)

    # --- 5. Commerce & subscriptions ---
    _heading(doc, "5. Commerce, wallets, and subscriptions", 1)
    _heading(doc, "5.1 Course checkout", 2)
    for line in (
        "Public checkout under /course/{id}/checkout with quote, coupon, wallet, and gateway flows.",
        "Gateways: Kashier callback route; Fawaterak prepare / methods / pay / return URLs (return may carry subscription or order correlation query params).",
        "Services coordinate pricing, wallet deduction, and order finalization (search app/Services for Checkout, Wallet, Coupon).",
    ):
        _bullet(doc, line)

    _heading(doc, "5.2 Teacher subscription (SaaS plans)", 2)
    for line in (
        "Public pricing pages and checkout for teacher_starter / teacher_pro style plans.",
        "Subscription requests, admin approval, online vs bank vs wallet semantics; activation services after gateway success.",
        "Student-facing PRO navigation is driven by active subscription rows and feature flags resolved on User.",
    ):
        _bullet(doc, line)

    _heading(doc, "5.3 Student subscription feature keys", 2)
    doc.add_paragraph(
        "config/student_subscription_features.php maps feature_key → UI route, icons, and labels. "
        "Keys include: library_access, ai_tools, classroom_access, support, teacher_profile, "
        "visible_to_academies, can_apply_opportunities, full_ai_suite, teacher_evaluation, "
        "recommended_to_academies, priority_opportunities, direct_support. "
        "Admin student-control screens consume the same config for auditing who has which feature."
    )

    _page_break(doc)

    # --- 6. Live & classroom ---
    _heading(doc, "6. Live sessions, recordings, and virtual classrooms", 1)
    for line in (
        "LiveSession model supports scheduled and live states; instructors drive sessions; students consume listings and join flows.",
        "Recording pipeline: webhooks/API (e.g. LiveRecordingWebhookController), admin review of live-recordings and classroom-recordings.",
        "Classroom: join codes under /classroom/join/{code}, heartbeat and leave endpoints, share-annotation channel, student-hosted meetings under student.classroom.* with recording upload (presigned/direct), audio tracks, AI-generated meeting reports (n8n-assisted updates for reports).",
        "Integration endpoints under /api/... for automation (n8n live session reports, classroom meeting reports).",
    ):
        _bullet(doc, line)

    _page_break(doc)

    # --- 7. Community & portfolio ---
    _heading(doc, "7. Community, portfolio, and growth", 1)
    for line in (
        "Community contributors: datasets, models (Model Zoo style pages), competitions — separate Community controllers and views.",
        "Public portfolio directory and per-teacher portfolio marketing pages.",
        "Teacher portfolio profile workflow: pending / approved / rejected with admin marketing review routes referenced from subscription feature config.",
        "Referral and loyalty program entities for growth tracking.",
    ):
        _bullet(doc, line)

    _page_break(doc)

    # --- Integrations snapshot ---
    _heading(doc, "7.1 Integrations and external systems", 2)
    _two_col_table(
        doc,
        [
            ("Fawaterak", "Card/online payments for course checkout and subscription checkout; return URL handling with correlation parameters."),
            ("Kashier", "Alternate payment gateway with callback route."),
            ("AWS S3", "Object storage for uploads, recordings, and scalable assets."),
            ("n8n / automation", "HTTP endpoints for updating live session and classroom meeting reports from workflows."),
            ("Google OAuth", "Social login path via Laravel Socialite."),
            ("WhatsApp", "WhatsAppHelper and admin WhatsApp settings for messaging integrations."),
            ("Video providers", "Admin-configurable providers for hosted lecture video delivery."),
        ],
        ("Integration", "Role"),
    )

    _page_break(doc)

    # --- 8. Data model ---
    _heading(doc, "8. Data model (representative entities)", 1)
    doc.add_paragraph(
        "The schema is broad. The following table lists core aggregates; refer to database/migrations "
        "and app/Models for the authoritative list."
    )
    _two_col_table(
        doc,
        [
            ("User", "Central identity; role; employee fields; 2FA; referral; portfolio columns."),
            ("Course / AdvancedCourse / Lesson / Lecture", "Content hierarchy, video providers, progress, materials."),
            ("Order / Payment / Invoice / Transaction / Wallet / WalletTransaction", "Monetary state machines and audit."),
            ("Subscription / SubscriptionRequest / Package", "Recurring access and feature entitlements."),
            ("Exam / ExamAttempt / QuestionBank", "Assessment domain."),
            ("Assignment / AssignmentSubmission", "Async coursework."),
            ("Certificate", "Completion credentials with public verification routes."),
            ("ClassroomMeeting / ClassroomMeetingParticipant / LiveSession / LiveRecording", "Synchronous teaching."),
            ("SupportTicket / SupportTicketReply", "Helpdesk for students/teachers."),
            ("EmployeeJob / EmployeeTask / LeaveRequest / Hr*", "Internal workforce domain."),
        ],
        ("Entity group", "Role in the platform"),
    )

    _page_break(doc)

    # --- 9. Developer guide ---
    _heading(doc, "9. Developer onboarding", 1)
    steps = doc.add_paragraph()
    for i, s in enumerate(
        [
            "Clone the repository and copy .env from .env.example; set APP_URL, database, mail, S3, and gateway keys.",
            "composer install && php artisan key:generate",
            "php artisan migrate --seed (use project-specific seeder guidance for roles/permissions).",
            "npm install && npm run build (if front-end assets are versioned via Vite/mix per project).",
            "Run php artisan serve or configure Apache/XAMPP virtual host to the public/ directory.",
            "Use route:list and IDE search on route('...') names to jump from UI to controllers.",
        ],
        1,
    ):
        steps.add_run(f"{i}. {s}\n")

    _heading(doc, "9.1 Where to change common behaviors", 2)
    _two_col_table(
        doc,
        [
            ("New public page", "routes/web.php + Public\\*Controller + resources/views/public/..."),
            ("New admin module", "routes/web.php admin group + Admin\\*Controller + resources/views/admin/... + permissions"),
            ("Student navigation", "resources/views/layouts/student-sidebar.blade.php + middleware on routes"),
            ("Subscription gating", "User subscription helpers + config/student_subscription_features.php + Subscription model JSON features"),
            ("Payment return URLs", "CheckoutController / SubscriptionCheckoutController + gateway dashboard"),
        ],
        ("Task", "Typical locations"),
    )

    _heading(doc, "9.2 Conventions", 2)
    for line in (
        "Prefer services for multi-step transactions (payments, activations) to keep controllers thin.",
        "Mirror permission names between seeders, middleware, and Blade directives.",
        "Bilingual UI: many strings use Laravel __() translations; documentation for product copy may live in lang/.",
    ):
        _bullet(doc, line)

    _page_break(doc)
    _add_engineering_deep_dive(doc)
    _add_auto_generated_inventory(doc)

    # --- Appendix (final quick reference) ---
    _heading(doc, "37. Appendix (quick reference)", 1)
    doc.add_paragraph("Quick reference for navigating routes/web.php and controller namespaces.")
    _heading(doc, "37.1 Route prefix map", 2)
    p = doc.add_paragraph()
    _mono_run(
        p,
        "/  (home, marketing)\n"
        "/course/*  (catalog + checkout)\n"
        "/dashboard  (persona router)\n"
        "/admin/*  (operations)\n"
        "/instructor/*  (authoring)\n"
        "/student/*  (some student routes; others at root with middleware)\n"
        "/employee/*  (internal desks)\n"
        "/classroom/join/*  (guest-friendly join)\n"
        "/api/*  (webhooks / n8n / recordings)\n",
        9,
    )

    _heading(doc, "37.2 Controller namespace index (app/Http/Controllers)", 2)
    doc.add_paragraph(
        "High-level map for repository navigation (non-exhaustive; reflects typical Laravel grouping):"
    )
    p = doc.add_paragraph()
    _mono_run(
        p,
        "Admin/          — platform governance, catalog, finance, HR, live ops, community moderation\n"
        "Instructor/     — teacher dashboards: courses, lessons, exams, live, calendar, payouts\n"
        "Student/        — enrolled learner tools: courses, wallet, subscription, classroom, AI usages\n"
        "Employee/       — internal desks (sales, HR, recruitment, agreements, leave)\n"
        "Public/         — marketing site, catalog, checkout, certificates verify, portfolio\n"
        "Community/      — contributor flows and community pages\n"
        "Auth/           — login, password, two-factor challenge\n"
        "Api/            — machine-facing webhooks and n8n callbacks\n",
        8,
    )

    _heading(doc, "37.3 Disclaimer", 2)
    doc.add_paragraph(
        "This document is generated to onboard engineers to the Muallimx monolith: it summarizes routes, "
        "middleware, services, and domains as reflected in the repository at generation time. It is not "
        "a substitute for php artisan route:list, migrations, or automated tests. Regenerate after major "
        "refactors and attach route:list exports for audit trails."
    )

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate Muallimx Word documentation.")
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=None,
        help="Output .docx path (default: Muallimx_Platform_Documentation_<date>.docx in cwd)",
    )
    args = parser.parse_args()

    cwd = Path.cwd()
    repo_hint = str(cwd)

    out = args.output
    if out is None:
        stamp = dt.datetime.now(timezone.utc).strftime("%Y%m%d")
        out = cwd / f"Muallimx_Platform_Documentation_{stamp}.docx"

    out = out.resolve()
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = build_document(repo_hint=repo_hint)
    doc.save(str(out))
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
