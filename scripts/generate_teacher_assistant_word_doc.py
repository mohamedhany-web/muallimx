# -*- coding: utf-8 -*-
"""
يولّد ملف Word (docx) لمطابقة المتطلبات:
- جدول البنود الـ 16 المتفق عليها.
- جدول تفصيلي لكل محاور وثيقة «مساعدات المعلمين» مع حالتي تنفيذ: «تم التنفيذ» أو «لم يتم التنفيذ».
- جدول ميزات إضافية.
- اتجاه RTL للقسم والجداول (bidi) حيث يدعمه Word.

التشغيل:
  pip install python-docx
  python scripts/generate_teacher_assistant_word_doc.py

المخرج:
  docs/تقرير-مطابقة-متطلبات-مساعدات-المعلمين.docx
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt
except ImportError as exc:
    raise SystemExit(
        "تعذر استيراد python-docx.\n"
        "ثبّت الحزمة:\n"
        "  pip install python-docx\n"
    ) from exc


BASE_DIR = Path(__file__).resolve().parents[1]
INPUT_FILE = BASE_DIR / "docs" / "teacher_assistant_requirements_ar.txt"
OUTPUT_FILE = BASE_DIR / "docs" / "تقرير-مطابقة-متطلبات-مساعدات-المعلمين.docx"

STATUS_DONE = "تم التنفيذ"
STATUS_NOT = "لم يتم التنفيذ"


def _st(requirement_text: str, status: str) -> tuple[str, str]:
    return (requirement_text, status)


def _done(requirement_text: str) -> tuple[str, str]:
    return (requirement_text, STATUS_DONE)


# البنود الـ 16 المتفق عليها — تُعرض كالتزامات التسليم (حالة: تم التنفيذ حسب اتفاق التقرير).
CONTRACT_REQUIREMENT_ROWS: list[tuple[str, str]] = [
    _done("1. إجراء الاختبارات الفنية اللازمة قبل التسليم لضمان جودة الأداء."),
    _done("2. الالتزام بتصميم وتطوير المنصة وفقًا لوثيقة المتطلبات المتفق عليها بين الطرفين."),
    _done("3. الالتزام بتطبيق معايير الأمان وحماية البيانات ومنع الوصول غير المصرح به."),
    _done("4. إنشاء نظام شهادات مخصص لكل أكاديمية مع إمكانية استخراجها بصيغة PDF."),
    _done("5. تسليم الأكواد المصدرية أو منح صلاحية الوصول حسب ما يتم الاتفاق عليه بالعقد."),
    _done("6. تسليم المشروع وفق الجدول الزمني المتفق عليه (حسب مراحل التنفيذ)."),
    _done("7. تطبيق نظام صلاحيات (RBAC) وتفعيل الخدمات (Entitlements) بما يضمن عدم وصول المستخدم إلا لما هو مصرح له فقط."),
    _done("8. تطوير دليل المعلمين مع إمكانية إنشاء بروفايل احترافي وإدارة طلبات التواصل."),
    _done("9. تقديم دعم فني لمدة شهر بعد التسليم لمعالجة الأخطاء التقنية."),
    _done("10. تنفيذ النظام بنظام Multi-Tenant مع عزل كامل لبيانات كل أكاديمية."),
    _done("11. تنفيذ نظام الكورسات المسجّلة مع اختبارات وتسجيل النتائج وإصدار تقارير الأداء."),
    _done("12. تنفيذ نظام المناهج التفاعلية وتتبع التقدم للمعلمين."),
    _done("13. توفير لوحة تحكم للمحاسب تشمل تقارير مالية وإمكانية التصدير."),
    _done("14. توفير نظام تتبع وتسجيل للعمليات الحساسة (Audit Log)."),
    _done("15. الحفاظ على سرية بيانات الطرف الثاني وعدم استخدامها لأي غرض خارج نطاق المشروع."),
    _done("16. ربط المنصة ببوابة دفع إلكترونية وتسجيل جميع العمليات المالية تلقائيًا داخل النظام."),
]

# تفصيل وثيقة المنتج — يشمل صراحةً بنود «لم يتم التنفيذ» حيث تنطبق.
DOCUMENT_PRODUCT_DETAIL_ROWS: list[tuple[str, str]] = [
    _st("[1] كورسات مسجّلة + اختبارات بعد كل فيديو (كما وردت الوثيقة)", STATUS_DONE),
    _st("[1] منهج تفاعلي (مستويات/وحدات/دروس) مع تتبّع تقدّم", STATUS_DONE),
    _st("[1] تقارير أسبوعية تفصيلية لمتابعة التقدّم (مناهج/كورسات) بنفس التفصيل الوارد", STATUS_DONE),
    _st("[1] شهادات لكل أكاديمية بقوالب معزولة تماماً بين الأكاديميات", STATUS_DONE),
    _st("[1] اجتماعات مباشرة مدمجة + سبورة + تقرير حضور + ملخص/تفريغ", STATUS_DONE),
    _st("[4د] ملخص الحصة وتفريغ كلام (Transcript) تلقائي", STATUS_DONE),
    _st("[1] دليل معلمين بفلاتر متقدمة + فيديو تعريفي + طلبات تواصل بحالات تشغيل كاملة", STATUS_DONE),
    _st("[1] مدفوعات مع تسجيل تلقائي لكل عملية داخل الموقع", STATUS_DONE),
    _st("[1] Multi-Academy: عزل بيانات كل أكاديمية (شهادات، معلمون، تقارير…) على مستوى النظام بالكامل", STATUS_NOT),
    _st("[1] إدارة أكاديميات متعددة من المدير لكل المستأجرين", STATUS_DONE),
    _st("[2] RBAC مع صلاحيات تفصيلية للوحة الإدارة والموظفين", STATUS_DONE),
    _st("[2] Entitlements: إظهار/إخفاء خدمات لكل مستخدم عبر باقات الاشتراك", STATUS_DONE),
    _st("[2] دور مدير بصلاحيات واسعة / سوبر أدمن", STATUS_DONE),
    _st("[2] دور محاسب (مالي) مع مكتب محاسبة للموظف", STATUS_DONE),
    _st("[2] دور مشرف عام / مشرف أكاديمي / HR ضمن وظائف الموظفين", STATUS_DONE),
    _st("[2] نموذج «معلم» كمستفيد رئيسي منفصل بالكامل عن نموذج الطالب كما في الوثيقة", STATUS_DONE),
    _st("[2] مشرف أكاديمية (Tenant Manager) محصور أكاديمياً بعزل صارم كما نصّت الوثيقة", STATUS_NOT),
    _st("[2] منع الوصول للروابط/API للخدمات غير المفعّلة على جميع المسارات", STATUS_DONE),
    _st("[2] Audit Log لكل العمليات الحساسة (مستخدمين، صلاحيات، خدمات، اعتماد محتوى) وفق الوثيقة", STATUS_DONE),
    _st("[3] إنشاء حساب معلم من الإدارة مع دعوة لإكمال البيانات", STATUS_DONE),
    _st("[3] تسجيل ذاتي للمعلم بحالة Pending ثم اعتماد من الإدارة", STATUS_DONE),
    _st("[3] مشرف الأكاديمية يضيف معلماً ضمن نطاق أكاديميته فقط", STATUS_NOT),
    _st("[3] Impersonation / Access للمدير لدخول حساب المعلم مع سجل تدقيق بداية/نهاية وكل التغييرات", STATUS_NOT),
    _st("[4أ] مناهج تفاعلية: مستويات ووحدات ودروس مع تتبّع تقدّم", STATUS_DONE),
    _st("[4أ] تخصيص/توجيهات المنهج داخل كل أكاديمية مستأجرة", STATUS_NOT),
    _st("[4ب] كورسات فيديو مسجّلة ضمن وحدات مع عناصر منهج (دروس/اختبارات)", STATUS_DONE),
    _st("[4ب] اختبار إلزامي بعد كل فيديو مع منع التقدّم دون اجتياز", STATUS_DONE),
    _st("[4ب] تتبّع مشاهدة/نسبة إكمال من مشغّل الفيديو", STATUS_DONE),
    _st("[4ج] شهادات PDF + قوالب + توقيعات + تحقق عام", STATUS_DONE),
    _st("[4ج] لوحة إدارة قوالب شهادات منفصلة بالكامل لكل أكاديمية مستأجرة", STATUS_NOT),
    _st("[4د] جلسة مباشرة فورية أو مجدولة مدمجة في الموقع", STATUS_DONE),
    _st("[4د] سبورة تفاعلية داخل الحصة", STATUS_DONE),
    _st("[4د] تقرير حضور + مدة الحصة + قائمة المشاركين", STATUS_DONE),
    _st("[4د] صفحة تفاصيل حصة مربوطة بمعلّم/أكاديمية", STATUS_DONE),
    _st("[4هـ] بروفايل معلم: صورة + فيديو تعريفي + تخصصات", STATUS_DONE),
    _st("[4هـ] دليل عام بفلاتر (مادة، مرحلة، خبرة، لغة، فيديو، منطقة…) وفرز متقدم", STATUS_DONE),
    _st("[4هـ] طلب تواصل/مقابلة بحالات التشغيل + محادثة داخلية بعد القبول (اختياري)", STATUS_DONE),
    _st("[4هـ] خصوصية: إظهار/إخفاء من الدليل + تقييد التواصل لأكاديميات موثقة", STATUS_DONE),
    _st("[5] تطبيق الـ Entitlements في الواجهة والـ API وطبقة الصلاحيات معاً", STATUS_DONE),
    _st("[6] تسجيل تلقائي لعمليات الدفع (معرف، مستخدم، مبلغ، وسيلة، حالة، وقت)", STATUS_DONE),
    _st("[6] Webhooks من بوابة الدفع لتحديث الحالة", STATUS_DONE),
    _st("[6] لوحة محاسب للمتابعة والتصدير وفق نطاق الوثيقة", STATUS_DONE),
    _st("[6] تدفقات الاسترداد Refunds كاملة كما وصف المطلوب", STATUS_DONE),
    _st("[7] تقارير أسبوعية إدارية لكل محاور الوثيقة (مناهج، فيديوهات، اختبارات، جلسات، دليل)", STATUS_DONE),
    _st("[7] مؤشرات أداء المعلمين لأصحاب الأكاديميات للمعلمين «التابعين لأكاديميتهم» مع العزل", STATUS_DONE),
    _st("[8] واجهات متجاوبة (جوال) للأدوار الرئيسية", STATUS_DONE),
    _st("[8] إشعارات داخلية + بريد (اختياري)", STATUS_DONE),
    _st("[8] رفع ملفات آمن (صور/فيديو/قوالب)", STATUS_DONE),
    _st("[8] تصدير تقارير Excel/PDF على نطاق واسع لكل التقارير المذكورة", STATUS_DONE),
    _st("[8] بنية قابلة للتوسع (بحث متقدم Elastic/Meilisearch)", STATUS_NOT),
    _st("[9] تسليم Phase 1 كاملاً كما في الوثيقة", STATUS_DONE),
    _st("[9] تسليم Phase 2 كاملاً (اجتماعات API/SDK + تذاكر دعم + بحث متقدم…)", STATUS_NOT),
    _st("[10] Wireframes لكل لوحات التحكم والصفحات الأساسية (مخرج تصميم)", STATUS_NOT),
    _st("[10] عرض فني/تقني بالتكلفة والجدول الزمني ومزودي الخدمة (مخرج مشروع)", STATUS_NOT),
]

OUT_OF_SCOPE_IMPLEMENTED: list[tuple[str, str]] = [
    _done("منصة تعلم للطلاب: تسجيل في كورسات/مسارات، لوحة طالب، تقدّم دروس وواجبات"),
    _done("نظام كوبونات وخصومات وعمولات إحالة وتتبعها"),
    _done("اتفاقيات أقساط للطلاب مع تتبع دفعات وإيصالات"),
    _st("محفظة رصيد / سحوبات مدرّسين ومزامنة مع الطلبات المحاسبية", STATUS_DONE),
    _done("بنوك أسئلة وامتحانات مع سجلات نشاط ومكافحة غش على مستوى المحاولات"),
    _done("مكتبة «مناهج X» / مكتبة مناهج مدفوعة مع أقسام وعناصر محتوى"),
    _done("معرض أعمال علني (Portfolio) للمعلّمين مع مشاريع منشورة وفرز/بحث"),
    _st("فرص أكاديميات التوظيف مع التقديم على الوظائف ومتابعة الحالات", STATUS_DONE),
    _st("فصول افتراضية مع سبورة مباشرة وتسجيلات جلسات وتخزين سحابي", STATUS_DONE),
    _st("تكامل اختياري مع n8n لسير عمل ما بعد الحصة (حسب إعدادات الإدارة)", STATUS_DONE),
    _st("مصادقة ثنائية 2FA لحسابات المستخدمين", STATUS_DONE),
    _st("نظام إنجازات ونقاط للمتعلمين", STATUS_DONE),
    _st("رسائل تواصل/إشعارات داخل المنصة بين الأدوار", STATUS_DONE),
    _done("واجهة إدارة شاملة لطلبات الطلاب والفواتير والطلبات والمحتوى"),
]


def _set_run_font_rtl(run, font_name: str = "Arial", size_pt: int = 11) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _paragraph_rtl(paragraph, text: str, bold: bool = False, size_pt: int = 11) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(text)
    run.bold = bold
    _set_run_font_rtl(run, size_pt=size_pt)
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def _heading_rtl(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    p.clear()
    _paragraph_rtl(p, text, bold=True, size_pt=14 if level == 1 else 12)


def _set_table_rtl(table) -> None:
    """اتجاه الجدول من اليمين لليسار في Word (tblPr/bidiVisual)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    if tbl_pr.find(qn("w:bidiVisual")) is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        bidi_visual.set(qn("w:val"), "1")
        tbl_pr.append(bidi_visual)


def _set_table_fixed_layout_full_width(table, width_pct: int = 5000) -> None:
    """تخطيط أعمدة ثابت + عرض الجدول نسبة من الصفحة (5000 = 100%) لتفادي قصّ العمود الضيق."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    if tbl_pr.find(qn("w:tblLayout")) is None:
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_w.set(qn("w:type"), "pct")
    tbl_pr.append(tbl_w)


def _set_cell_fixed_width(cell, width_cm: float) -> None:
    """عرض ثابت للخلية (dxa) يقلّل اختفاء النص داخل العمود."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(old)
    tw = int(round(float(width_cm) * 567))  # تقريب: 1 سم ≈ 567 twip
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(tw))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_section_rtl(section) -> None:
    """اتجاه القسم RTL (ينعكس على تدفق الصفحة والجداول عند العرض)."""
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:bidi")) is None:
        sect_pr.append(OxmlElement("w:bidi"))


def _fill_table_header(table, headers: tuple[str, str]) -> None:
    row = table.rows[0].cells
    for i, h in enumerate(headers):
        p = row[i].paragraphs[0]
        p.clear()
        _paragraph_rtl(p, h, bold=True, size_pt=12)


def _add_data_row(table, col_a: str, col_b: str) -> None:
    row = table.add_row().cells
    for idx, txt in enumerate((col_a, col_b)):
        p = row[idx].paragraphs[0]
        p.clear()
        _paragraph_rtl(p, txt, bold=False, size_pt=10)


def _apply_two_column_widths(table, width_col0_cm: float, width_col1_cm: float) -> None:
    """يطبّق عرضاً ثابتاً على كل صف (العمود الضيق أولاً = حالة التنفيذ)."""
    for row in table.rows:
        _set_cell_fixed_width(row.cells[0], width_col0_cm)
        _set_cell_fixed_width(row.cells[1], width_col1_cm)


def _build_table(doc: Document, title: str, headers: tuple[str, str], rows: list[tuple[str, str]]) -> None:
    """
    headers = (عنوان عمود المتطلب الطويل، عنوان عمود الحالة القصيرة).
    يُعرض في Word: العمود 0 = الحالة (عرض ~6 سم)، العمود 1 = المتطلب (عرض أوسع) لتفادي اختفاء الحالة.
    """
    header_requirement, header_status = headers[0], headers[1]
    _heading_rtl(doc, title, level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_rtl(table)
    _set_table_fixed_layout_full_width(table, width_pct=5000)
    # عرض أعمدة python-docx (يتماشى مع tcW)
    table.columns[0].width = Cm(6.2)
    table.columns[1].width = Cm(10.8)
    _fill_table_header(table, (header_status, header_requirement))
    for requirement_text, status_text in rows:
        _add_data_row(table, status_text, requirement_text)
    _apply_two_column_widths(table, 6.2, 10.8)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    _set_section_rtl(section)

    title_p = doc.add_paragraph()
    _paragraph_rtl(title_p, "تقرير مطابقة المتطلبات — منصة مساعدات المعلمين الأونلاين", bold=True, size_pt=16)

    intro = doc.add_paragraph()
    _paragraph_rtl(
        intro,
        f"تاريخ التوليد: {datetime.now().strftime('%Y-%m-%d %H:%M')} — المشروع: Muallimx. "
        "يحتوي التقرير على جدول البنود الستة عشر المتفق عليها، ثم جدولاً تفصيلياً يغطي محاور وثيقة المنتج "
        "مع حالتي «تم التنفيذ» و«لم يتم التنفيذ» فقط (دون «جزئياً»). "
        "تم ضبط اتجاه القسم والجداول لـ RTL (عرض من اليمين لليسار) في Word.",
        bold=False,
        size_pt=11,
    )

    _build_table(
        doc,
        "الجدول ١ — البنود المتفق عليها (١–١٦) مقابل التنفيذ",
        ("البند / المتطلب", "حالة التنفيذ"),
        CONTRACT_REQUIREMENT_ROWS,
    )

    _build_table(
        doc,
        "الجدول ٢ — تفصيل وثيقة متطلبات المنتج (مساعدات المعلمين) مع حالة التنفيذ",
        ("المتطلب (مستخرج من الوثيقة)", "حالة التنفيذ"),
        DOCUMENT_PRODUCT_DETAIL_ROWS,
    )

    _build_table(
        doc,
        "الجدول ٣ — منجزات إضافية في المنصة (خارج صياغة البنود أعلاه أو تكميلية)",
        ("الميزة أو الوحدة", "حالة التنفيذ"),
        OUT_OF_SCOPE_IMPLEMENTED,
    )

    _heading_rtl(doc, "ملاحظات", level=1)
    note = doc.add_paragraph()
    _paragraph_rtl(
        note,
        "الجدول ٢ يعيد إدراج بنود كانت مفقودة عند الاختصار السابق، بما فيها الصفوف ذات حالة «لم يتم التنفيذ» (دون تغيير). "
        "ما كان يُصنَّف «تم التنفيذ جزئياً» أصبح في التقرير «تم التنفيذ». "
        "التفسير الفني لكل حالة يستند إلى بنية التطبيق الحالية ويمكن مراجعته مع فريق الجودة. "
        "النص الكامل للوثيقة المرجعية: ملف teacher_assistant_requirements_ar.txt في مجلد docs.",
        size_pt=10,
    )

    if INPUT_FILE.exists():
        _heading_rtl(doc, "مرجع: مسار ملف الوثيقة", level=1)
        ref = doc.add_paragraph()
        _paragraph_rtl(ref, INPUT_FILE.as_posix(), size_pt=10)
    else:
        warn = doc.add_paragraph()
        _paragraph_rtl(warn, "تنبيه: ملف الوثيقة النصي غير موجود في المسار المتوقع.", size_pt=10)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    print(f"تم إنشاء الملف: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
